import os
import logging

from chatbot.models import CompanyBot
from chatbot.models.story_vernacular_model import StoryVernacular
from chatbot.utils.S3.s3_service import upload_file_to_s3
from chatbot.utils.gotenberg_utils import generate_pdf_with_gotenberg
from chatbot.models.enums import MediaTypeChoices

logger = logging.getLogger('django')


def _add_hyperlink(paragraph, url, text):
    """Insert a clickable hyperlink run into a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r_id = paragraph.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run_elem = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rpr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.append(underline)

    run_elem.append(rpr)

    t = OxmlElement('w:t')
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def create_pdf_from_text(text_content, company_bot_id) -> bytes:
    """
    Create a PDF file from text content using Gotenberg HTML-to-PDF service.
    """
    try:
        # Convert text to formatted HTML
        html_content = text_to_html(text_content, company_bot_id)
        
        # Use Gotenberg to convert HTML to PDF
        pdf_content = generate_pdf_with_gotenberg(html_content)
        
        if not pdf_content:
            raise Exception("Gotenberg failed to generate PDF")
        
        logger.info(f"Successfully created PDF with {len(text_content)} characters")
        
        return pdf_content
        
    except Exception as e:
        logger.error(f"Error creating PDF from text: {e}", exc_info=True)
        raise


def text_to_html(text_content, company_bot_id) -> str:
    """
    Convert Markdown text to styled HTML for PDF generation.
    """
    import markdown
    import re

    max_char_per_page = 1500
    try:
        print(f"company_bot_id: {company_bot_id}")
        logger.info(f"company_bot_id: {company_bot_id}")
        company_bot = CompanyBot.objects.filter(id=company_bot_id).first()
        print(f"company_bot: {company_bot}")
        logger.info(f"company_bot: {company_bot}")
        if company_bot:
            story_vernacular = StoryVernacular.objects.filter(
                company_bot=company_bot, language='en'
            ).first()
            print(f"story_vernacular: {story_vernacular}")
            logger.info(f"story_vernacular: {story_vernacular}")
            if story_vernacular and story_vernacular.translation_json:
                logger.info(f"story_vernacular translation_json: {story_vernacular.translation_json}")
                max_char_per_page = story_vernacular.translation_json.get(
                    'page_split_char_len', max_char_per_page
                )
        logger.info(f"Using max_char_per_page: {max_char_per_page}")
    except Exception as e:
        logger.info(f"Could not get max_char_per_page from StoryVernacular: {e}")

    # Markdown → HTML
    html = markdown.markdown(
        text_content,
        extensions=[
            "tables",
            "fenced_code",
            "sane_lists",
            "toc",
            "def_list"
        ]
    )

    # Regex to find block-level elements (with attributes allowed)
    block_pattern = re.compile(
        r'(<h[1-3][^>]*>.*?</h[1-3]>|'
        r'<p[^>]*>.*?</p>|'
        r'<ul[^>]*>.*?</ul>|'
        r'<ol[^>]*>.*?</ol>|'
        r'<table[^>]*>.*?</table>)',
        flags=re.DOTALL
    )

    pages = []
    current_page = ""
    char_count = 0
    last_index = 0

    def is_heading(block: str) -> bool:
        return block.lstrip().startswith("<h")

    def is_content_block(block: str) -> bool:
        return block.lstrip().startswith(("<p", "<ul", "<ol", "<table"))

    for match in block_pattern.finditer(html):
        start, end = match.span()

        # Preserve any content BEFORE this block
        prefix = html[last_index:start]
        if prefix:
            current_page += prefix
            char_count += len(prefix)

        block = match.group()
        current_page += block
        char_count += len(block)

        # Insert page break ONLY after content blocks
        if char_count >= max_char_per_page and is_content_block(block):
            pages.append(current_page)
            current_page = ""
            char_count = 0

        last_index = end

    # Append remaining tail content
    tail = html[last_index:]
    if tail:
        current_page += tail

    if current_page.strip():
        pages.append(current_page)

    # Wrap pages
    paginated_html = ""
    for page in pages:
        paginated_html += f"""
        <div class="page">
            {page}
        </div>
        """

    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: 'Open Sans', Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                margin: 2.2cm;
                color: #111;
            }}

            .page {{
                page-break-after: always;
            }}

            .page:last-child {{
                page-break-after: auto;
            }}

            h1 {{
                font-size: 18pt;
                text-align: center;
                margin: 0 0 1.5em 0;
                text-transform: uppercase;
                letter-spacing: 0.4px;
            }}

            h2 {{
                font-size: 14pt;
                margin-top: 1.8em;
                margin-bottom: 0.8em;
                font-weight: 600;
            }}

            h3 {{
                font-size: 12pt;
                margin-top: 1.4em;
                margin-bottom: 0.5em;
                font-weight: 600;
            }}

            p {{
                margin-bottom: 0.9em;
                text-align: left;
            }}

            ul, ol {{
                margin-left: 1.4em;
                margin-bottom: 1em;
            }}

            li {{
                margin-bottom: 0.3em;
            }}

            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 1.4em 0;
                font-size: 9.5pt;
            }}

            th, td {{
                border: 1px solid #555;
                padding: 6px 8px;
                vertical-align: top;
            }}

            th {{
                background-color: #e9ecef;
                font-weight: bold;
                text-align: center;
            }}
        </style>
    </head>
    <body>
        {paginated_html}
    </body>
    </html>
    """

    return html_template


def sanitize_filename(filename: str, extension: str = '.pdf') -> str:
    """Sanitize filename and enforce the given extension."""
    ext = extension if extension.startswith('.') else f'.{extension}'
    try:
        filename = os.path.basename(filename)
        name_without_ext = os.path.splitext(filename)[0]
        safe_name = "".join(c for c in name_without_ext if c.isalnum() or c in (' ', '-', '_'))
        safe_name = '_'.join(safe_name.split())
        if not safe_name:
            safe_name = "download"
        return f"{safe_name}{ext}"
    except Exception as e:
        logger.error(f"Error sanitizing filename: {e}")
        return f"download{ext}"


def render_template_to_pdf(
    *,
    flow_name: str,
    arguments: dict,
    company_bot_id: int,
    session_id: str,
    sources: list = None,
) -> dict:
    """
    Look up the PDFTemplate for the flow (by flow_route), render it with Jinja2,
    generate a PDF via Gotenberg, and upload to S3.

    Falls back to create_and_upload_file if no template is found.
    """
    from jinja2 import Template
    from chatbot.models.company_models import Flow, PDFTemplates
    from chatbot.models.chat_models import ChatSession

    try:
        print("flow_route: ", flow_name)
        flow = Flow.objects.filter(flow_route=flow_name).first() if flow_name else None
        print("Flow found: ", flow)
        pdf_template = PDFTemplates.objects.filter(flow=flow).first() if flow else None
        print("PDF Template: ", pdf_template)

        logger.info(f'[render_template_to_pdf] flow_name={flow_name!r} flow={flow} pdf_template={pdf_template}')

        if not pdf_template:
            logger.info(
                f'[render_template_to_pdf] No PDFTemplate found for flow_route={flow_name!r} — falling back to text-based PDF'
            )
            fallback_content = arguments.get('knowledge_content') or str(arguments)
            return create_and_upload_file(
                content=fallback_content,
                filename=arguments.get('filename', 'download.pdf'),
                company_bot_id=company_bot_id,
                session_id=session_id,
            )

        chat_session = ChatSession.objects.filter(session=session_id).first()
        profile = chat_session.profile if chat_session else None

        context = {
            'args': arguments,
            'constants': pdf_template.constants_json or {},
            'profile': profile,
            'sources': sources or arguments.get('sources') or [],
        }

        html_content = Template(pdf_template.template).render(**context)

        pdf_content = generate_pdf_with_gotenberg(html_content)
        if not pdf_content:
            raise Exception('Gotenberg failed to generate PDF from template')

        safe_filename = sanitize_filename(arguments.get('filename', 'download.pdf'), '.pdf')
        s3_key = upload_file_to_s3(
            file_name=safe_filename,
            file_content=pdf_content,
            content_type=MediaTypeChoices.PDF,
            project_id=None,
            folder_structure=f'chatbot/{company_bot_id}/{session_id}/',
        )

        if not s3_key:
            return {'success': False, 'error': 'Failed to upload PDF to S3'}

        return {'success': True, 'media_url': f'{os.getenv("S3_MEDIA_URL")}{s3_key}', 'file_name': safe_filename}

    except Exception as e:
        logger.error(f'[render_template_to_pdf] Error: {e}', exc_info=True)
        return {'success': False, 'error': str(e)}


def create_docx_from_args(
    *,
    arguments: dict,
    company_bot_id: int,
    session_id: str,
    sources: list = None,
) -> dict:
    """
    Generate a DOCX file directly from download_file tool call arguments (no template model).
    MIP documents get structured sections; knowledge documents get the raw content.
    """
    import io
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        is_mip = bool(arguments.get('goal') or arguments.get('action_plan'))
        safe_filename = sanitize_filename(arguments.get('filename', 'download.docx'), '.docx')

        doc = docx.Document()

        title_text = os.path.splitext(arguments.get('filename', 'Document'))[0].replace('_', ' ').replace('-', ' ')

        if is_mip:
            doc.add_heading(title_text, level=1)

            if arguments.get('goal'):
                doc.add_heading('Goal', level=2)
                doc.add_paragraph(arguments['goal'])

            if arguments.get('objective'):
                doc.add_heading('Objective', level=2)
                doc.add_paragraph(arguments['objective'])

            if arguments.get('duration'):
                doc.add_heading('Timeline', level=2)
                doc.add_paragraph(f"Duration: {arguments['duration']}")

            action_plan = arguments.get('action_plan') or []
            if action_plan:
                from docx.shared import Inches
                doc.add_heading('Action plan', level=2)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                header_cells = table.rows[0].cells
                header_cells[0].text = '#'
                header_cells[1].text = 'Action'
                header_cells[2].text = 'Week'
                table.columns[0].width = Inches(0.4)
                table.columns[1].width = Inches(5.0)
                table.columns[2].width = Inches(1.1)
                for i, step in enumerate(action_plan):
                    row = table.add_row().cells
                    row[0].text = str(i + 1)
                    row[1].text = step.get('action', '')
                    row[2].text = step.get('week', '')

            success_indicators = arguments.get('success_indicators') or []
            if success_indicators:
                doc.add_heading('Success indicators', level=2)
                for i, indicator in enumerate(success_indicators):
                    doc.add_paragraph(f"{i + 1}. {indicator}")

        else:
            doc.add_heading(title_text, level=1)
            content = arguments.get('knowledge_content', '')
            for para in content.split('\n\n'):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

        resolved_sources = sources or arguments.get('sources') or []
        if resolved_sources:
            doc.add_heading('References', level=2)
            for src in resolved_sources:
                src_title = src.get('title', '')
                src_url = src.get('url', '')
                para = doc.add_paragraph()
                if src_url:
                    _add_hyperlink(para, url=src_url, text=src_title or src_url)
                else:
                    para.add_run(src_title)

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        s3_key = upload_file_to_s3(
            file_name=safe_filename,
            file_content=docx_bytes,
            content_type=MediaTypeChoices.DOCX,
            project_id=None,
            folder_structure=f'chatbot/{company_bot_id}/{session_id}/',
        )

        if not s3_key:
            return {'success': False, 'error': 'Failed to upload DOCX to S3'}

        return {'success': True, 'media_url': f'{os.getenv("S3_MEDIA_URL")}{s3_key}', 'file_name': safe_filename}

    except Exception as e:
        logger.error(f'[create_docx_from_args] Error: {e}', exc_info=True)
        return {'success': False, 'error': str(e)}


def create_and_upload_file(
    *,
    content: str,
    filename: str,
    company_bot_id: int,
    session_id: str
) -> dict:
    """
    Create a PDF file from content and upload it to S3.
    """
    try:
        logger.info(f"Creating file for session {session_id}, company_bot {company_bot_id}")
        logger.info(f"Original filename: {filename}, content length: {len(content)} chars")
        
        # Sanitize filename and ensure .pdf extension
        safe_filename = sanitize_filename(filename)
        logger.info(f"Sanitized filename: {safe_filename}")
        
        # Create PDF from content using Gotenberg
        pdf_content = create_pdf_from_text(content, company_bot_id)
        
        logger.info(f"PDF created successfully, size: {len(pdf_content)} bytes")
        
        folder_structure = f"chatbot/{company_bot_id}/{session_id}/"
        
        # Upload to S3
        s3_key = upload_file_to_s3(
            file_name=safe_filename,
            file_content=pdf_content,
            content_type=MediaTypeChoices.PDF,
            project_id=None,
            folder_structure=folder_structure
        )
        
        if not s3_key:
            logger.error("Failed to upload file to S3")
            return {
                'success': False,
                'error': 'Failed to upload file to S3'
            }
        
        # Construct media URL
        base = os.getenv("S3_MEDIA_URL")
        media_url = f"{base}{s3_key}"
        
        logger.info(f"File uploaded successfully: {media_url}")
        
        return {
            'success': True,
            'media_url': media_url,
            'file_name': safe_filename,
            's3_key': s3_key
        }
        
    except Exception as e:
        logger.error(f"Error creating and uploading file: {e}", exc_info=True)
        return {
            'success': False,
            'error': str(e)
        }
